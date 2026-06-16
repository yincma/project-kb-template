from __future__ import annotations

from pathlib import Path
from typing import Any

from kb.parsers.ocr import OCRProcessor
from kb.parsers.registry import ParsedDocument, ParsedSection
from kb.multimodal.extraction import new_visual_stats, register_visual_bytes


def parse_docx_file(path: Path, config: Any | None = None) -> ParsedDocument:
    warnings: list[str] = []
    try:
        from docx import Document
    except Exception as exc:
        return ParsedDocument(path=path, warnings=[f"Failed to parse {path}: python-docx is unavailable: {exc}"])

    ocr_cfg = _ocr_config(config)
    office_cfg = _office_config(config)
    ocr = OCRProcessor(_cfg_value(ocr_cfg, "engine", "rapidocr")) if _cfg_value(ocr_cfg, "enabled", True) else None
    extract_images = bool(_cfg_value(office_cfg, "extract_images", True))

    try:
        document = Document(str(path))
    except Exception as exc:
        return ParsedDocument(path=path, warnings=[f"Failed to parse {path}: {exc}"])

    sections: list[ParsedSection] = []
    visual_stats = new_visual_stats()
    current_heading: str | None = None
    current_parts: list[str] = []

    def flush() -> None:
        nonlocal current_parts
        text = "\n\n".join(part for part in current_parts if part.strip()).strip()
        if text:
            sections.append(
                ParsedSection(
                    text=text,
                    heading=current_heading,
                    parser_name="docx",
                    source_format=".docx",
                    extraction_method="text",
                    asset_type="document",
                )
            )
        current_parts = []

    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        if style_name.startswith("heading"):
            flush()
            current_heading = text
            current_parts.append(text)
        else:
            current_parts.append(text)

    for table in document.tables:
        table_text = _table_text(table)
        if table_text:
            current_parts.append(table_text)

    if extract_images and ocr:
        for rel_id, part in document.part.related_parts.items():
            content_type = getattr(part, "content_type", "")
            if not content_type.startswith("image/"):
                continue
            try:
                ocr_result = ocr.ocr_image(part.blob)
                if ocr_result.text.strip():
                    current_parts.append(f"Image OCR {rel_id}:\n{ocr_result.text.strip()}")
            except Exception as exc:
                warnings.append(f"OCR failed for {path} image {rel_id}: {exc}")

    if _multimodal_enabled(config):
        image_counter = 0
        context_text = "\n".join(current_parts)[-2000:]
        for rel_id, part in document.part.related_parts.items():
            content_type = getattr(part, "content_type", "")
            if not content_type.startswith("image/"):
                continue
            image_counter += 1
            try:
                ext = "." + content_type.split("/", 1)[1].replace("jpeg", "jpg")
                section = register_visual_bytes(
                    image_bytes=part.blob,
                    source_path=path,
                    config=config,
                    generated_from="embedded_image",
                    occurrence_index=image_counter,
                    context_title=current_heading,
                    nearby_text=context_text,
                    ext=ext,
                    stats=visual_stats,
                )
                if section:
                    sections.append(section)
            except Exception as exc:
                warnings.append(f"Multimodal image extraction failed for {path} image {rel_id}: {exc}")

    flush()

    if any("Image OCR" in section.text for section in sections):
        for section in sections:
            if "Image OCR" in section.text:
                section.ocr_used = True
                section.extraction_method = "text+ocr"

    warnings.extend(str(warning) for warning in visual_stats.get("warnings", []))
    warnings.extend(str(warning) for warning in visual_stats.get("limit_warnings", []))
    return ParsedDocument(path=path, sections=sections, warnings=warnings)


def _table_text(table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            rows.append("\t".join(cells))
    return "\n".join(rows)


def _ocr_config(config: Any | None) -> Any | None:
    return _cfg_value(_cfg_value(config, "parsing", None), "ocr", None)


def _office_config(config: Any | None) -> Any | None:
    return _cfg_value(_cfg_value(config, "parsing", None), "office", None)


def _multimodal_enabled(config: Any | None) -> bool:
    return bool(_cfg_value(_cfg_value(_cfg_value(config, "parsing", None), "multimodal", None), "enabled", False))


def _cfg_value(config: Any | None, name: str, default: Any = None) -> Any:
    if config is None:
        return default
    if isinstance(config, dict):
        return config.get(name, default)
    return getattr(config, name, default)
